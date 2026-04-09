#!/usr/bin/env -S uv run --script
# /// script
# requires-python = ">=3.12"
# dependencies = [
#   "faker>=20.0",
#   "click>=8.0",
#   "openpyxl>=3.1",
#   "python-docx>=1.0",
#   "reportlab>=4.0",
#   "Pillow>=10.0",
# ]
# ///
"""
Test Data Generator -- runs standalone with `uv run generate.py`
No manual pip install or venv needed.

Usage examples:
  # Explicit size
  uv run generate.py json  --rows 5000
  uv run generate.py csv   --rows 100000 --output big.csv
  uv run generate.py xlsx  --rows 50000  --schema products
  uv run generate.py docx  --pages 200
  uv run generate.py pdf   --pages 300   --output book.pdf
  uv run generate.py html  --sections 100
  uv run generate.py xml   --rows 5000
  uv run generate.py md    --sections 50
  uv run generate.py image --width 1920 --height 1080 --format jpg
  uv run generate.py lorem --words 10000
  uv run generate.py zip

  # Size presets (small / medium / large)
  uv run generate.py pdf   --size large
  uv run generate.py json  --size large --schema transactions

  # Batch -- one file of every format
  uv run generate.py batch --size small --output ./fixtures/

  # Reproducible output
  uv run generate.py json --rows 1000 --seed 42

  # Localized data
  uv run generate.py csv --rows 500 --locale fr_FR

  # Different schemas (for json / csv / xlsx)
  uv run generate.py json --rows 1000 --schema users
  uv run generate.py json --rows 1000 --schema products
  uv run generate.py json --rows 1000 --schema transactions
  uv run generate.py json --rows 1000 --schema logs
  uv run generate.py json --rows 1000 --schema articles
  uv run generate.py json --rows 1000 --schema lorem
"""

from __future__ import annotations

import csv as _csv
import json as _json
import sys
import xml.etree.ElementTree as ET
import zipfile
from datetime import datetime
from pathlib import Path

import click
from faker import Faker

# ---------------------------------------------------------------------------
# Size presets
# ---------------------------------------------------------------------------

SIZE_PRESETS: dict[str, dict] = {
    "small":  {"pages": 5,    "rows": 100,     "sections": 5,   "words": 500,    "width": 400,  "height": 300},
    "medium": {"pages": 50,   "rows": 5_000,   "sections": 20,  "words": 5_000,  "width": 800,  "height": 600},
    "large":  {"pages": 500,  "rows": 100_000, "sections": 100, "words": 50_000, "width": 1920, "height": 1080},
}

SIZE_CHOICES = click.Choice(list(SIZE_PRESETS), case_sensitive=False)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _make_faker(locale: str, seed: int | None) -> Faker:
    fake = Faker(locale, use_weighting=False)
    if seed is not None:
        Faker.seed(seed)
    return fake


def _progress(iterable, label: str, total: int):
    """Auto progress bar when stdout is a TTY, silent otherwise."""
    if sys.stdout.isatty():
        with click.progressbar(iterable, length=total, label=label) as bar:
            yield from bar
    else:
        yield from iterable


def _resolve(explicit: int | None, preset_key: str, size: str | None) -> int:
    """Return explicit value, or look up preset_key in the chosen size preset (default: medium)."""
    if explicit is not None:
        return explicit
    return SIZE_PRESETS[size or "medium"][preset_key]


def _file_size_label(path: Path) -> str:
    b = path.stat().st_size
    return f"{b / 1_048_576:.1f} MB" if b >= 1_048_576 else f"{b / 1024:.1f} KB"


# ---------------------------------------------------------------------------
# Record builders  (one dict per row)
# ---------------------------------------------------------------------------

SCHEMAS: dict[str, callable] = {}


def _schema(name: str):
    def decorator(fn):
        SCHEMAS[name] = fn
        return fn
    return decorator


@_schema("users")
def _users_record(fake: Faker, i: int) -> dict:
    return {
        "id": i + 1,
        "name": fake.name(),
        "email": fake.email(),
        "phone": fake.phone_number(),
        "company": fake.company(),
        "job": fake.job(),
        "address": fake.address().replace("\n", ", "),
        "country": fake.country(),
        "birthdate": fake.date_of_birth(minimum_age=18, maximum_age=80).isoformat(),
        "registered_at": fake.date_time_this_decade().isoformat(),
        "active": fake.boolean(),
        "score": round(fake.pyfloat(min_value=0, max_value=100, right_digits=2), 2),
    }


@_schema("products")
def _products_record(fake: Faker, i: int) -> dict:
    categories = ["Electronics", "Clothing", "Books", "Food", "Sports", "Home", "Toys"]
    return {
        "id": i + 1,
        "sku": fake.unique.bothify(text="??-######").upper(),
        "name": fake.catch_phrase(),
        "description": fake.paragraph(nb_sentences=3),
        "category": fake.random_element(categories),
        "brand": fake.company(),
        "price": round(fake.pyfloat(min_value=0.99, max_value=9999.99, right_digits=2), 2),
        "stock": fake.pyint(min_value=0, max_value=10000),
        "rating": round(fake.pyfloat(min_value=1, max_value=5, right_digits=1), 1),
        "tags": fake.words(nb=3),
        "created_at": fake.date_time_this_year().isoformat(),
    }


@_schema("transactions")
def _transactions_record(fake: Faker, i: int) -> dict:
    statuses = ["pending", "completed", "failed", "refunded"]
    methods = ["credit_card", "debit_card", "bank_transfer", "paypal", "crypto"]
    return {
        "id": fake.uuid4(),
        "seq": i + 1,
        "from_account": fake.bban(),
        "to_account": fake.bban(),
        "currency": fake.currency_code(),
        "amount": round(fake.pyfloat(min_value=1, max_value=100000, right_digits=2), 2),
        "method": fake.random_element(methods),
        "status": fake.random_element(statuses),
        "timestamp": fake.date_time_this_year().isoformat(),
        "description": fake.sentence(nb_words=6),
        "ip_address": fake.ipv4(),
    }


@_schema("logs")
def _logs_record(fake: Faker, i: int) -> dict:
    levels = ["DEBUG", "INFO", "INFO", "INFO", "WARNING", "ERROR", "CRITICAL"]
    methods = ["GET", "POST", "PUT", "DELETE", "PATCH"]
    return {
        "seq": i + 1,
        "timestamp": fake.date_time_this_month().isoformat(),
        "level": fake.random_element(levels),
        "service": fake.random_element(["api", "worker", "scheduler", "gateway", "auth"]),
        "request_id": fake.uuid4(),
        "method": fake.random_element(methods),
        "path": "/" + "/".join(fake.words(nb=fake.pyint(min_value=1, max_value=4))),
        "status_code": fake.random_element([200, 200, 200, 201, 400, 401, 403, 404, 500]),
        "duration_ms": fake.pyint(min_value=1, max_value=5000),
        "user_agent": fake.user_agent(),
        "ip": fake.ipv4(),
        "message": fake.sentence(),
    }


@_schema("articles")
def _articles_record(fake: Faker, i: int) -> dict:
    return {
        "id": i + 1,
        "title": fake.sentence(nb_words=8).rstrip("."),
        "slug": "-".join(fake.words(nb=5)),
        "author": fake.name(),
        "category": fake.word(),
        "tags": fake.words(nb=4),
        "summary": fake.paragraph(nb_sentences=2),
        "body": "\n\n".join(fake.paragraphs(nb=5)),
        "views": fake.pyint(min_value=0, max_value=1_000_000),
        "likes": fake.pyint(min_value=0, max_value=50_000),
        "published_at": fake.date_time_this_year().isoformat(),
        "updated_at": fake.date_time_this_month().isoformat(),
    }


@_schema("lorem")
def _lorem_record(fake: Faker, i: int) -> dict:
    return {
        "id": i + 1,
        "title": fake.sentence(nb_words=5).rstrip("."),
        "body": fake.paragraph(nb_sentences=4),
        "value": fake.pyint(),
        "flag": fake.boolean(),
    }


# ---------------------------------------------------------------------------
# File generators
# ---------------------------------------------------------------------------

def gen_json(fake: Faker, rows: int, output: Path, schema: str) -> None:
    """Streaming JSON write -- memory-efficient for very large files."""
    builder = SCHEMAS[schema]
    with output.open("w", encoding="utf-8") as f:
        f.write("[\n")
        for i in _progress(range(rows), label="Building JSON", total=rows):
            record = builder(fake, i)
            f.write("  " + _json.dumps(record, ensure_ascii=False, default=str))
            f.write(",\n" if i < rows - 1 else "\n")
        f.write("]\n")


def gen_csv(fake: Faker, rows: int, output: Path, schema: str) -> None:
    builder = SCHEMAS[schema]
    first = builder(fake, 0)

    def _flatten(record: dict) -> dict:
        return {k: (", ".join(str(x) for x in v) if isinstance(v, list) else v) for k, v in record.items()}

    with output.open("w", newline="", encoding="utf-8") as f:
        writer = _csv.DictWriter(f, fieldnames=list(first.keys()))
        writer.writeheader()
        writer.writerow(_flatten(first))
        for i in _progress(range(1, rows), label="Writing CSV", total=rows - 1):
            writer.writerow(_flatten(builder(fake, i)))


def gen_xlsx(fake: Faker, rows: int, output: Path, schema: str) -> None:
    from openpyxl import Workbook

    builder = SCHEMAS[schema]
    first = builder(fake, 0)
    headers = list(first.keys())

    def _flatten_row(record: dict) -> list:
        return [", ".join(str(x) for x in v) if isinstance(v, list) else v for v in record.values()]

    wb = Workbook(write_only=True)
    ws = wb.create_sheet("Data")
    ws.append(headers)
    ws.append(_flatten_row(first))
    for i in _progress(range(1, rows), label="Writing XLSX", total=rows - 1):
        ws.append(_flatten_row(builder(fake, i)))
    wb.save(str(output))


def gen_docx(fake: Faker, pages: int, output: Path) -> None:
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    doc.styles["Heading 1"].font.size = Pt(16)
    doc.styles["Normal"].font.size = Pt(11)
    doc.add_heading(fake.company(), level=0)

    for i in _progress(range(pages), label="Building DOCX", total=pages):
        doc.add_heading(f"Section {i + 1}: {fake.catch_phrase()}", level=1)
        doc.add_paragraph(fake.paragraph(nb_sentences=6))
        doc.add_paragraph(fake.paragraph(nb_sentences=4))
        if i % 5 == 0:
            table = doc.add_table(rows=1, cols=3)
            table.style = "Table Grid"
            hdr = table.rows[0].cells
            hdr[0].text, hdr[1].text, hdr[2].text = "Name", "Email", "Note"
            for _ in range(4):
                row = table.add_row().cells
                row[0].text = fake.name()
                row[1].text = fake.email()
                row[2].text = fake.sentence(nb_words=5)
        doc.add_page_break()

    doc.save(str(output))


def gen_pdf(fake: Faker, pages: int, output: Path) -> None:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.units import cm
    from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

    story = []
    styles = getSampleStyleSheet()

    for i in _progress(range(pages), label="Building PDF", total=pages):
        story.append(Paragraph(f"Chapter {i + 1}: {fake.catch_phrase()}", styles["Heading1"]))
        story.append(Spacer(1, 0.3 * cm))
        for _ in range(3):
            story.append(Paragraph(fake.paragraph(nb_sentences=5), styles["Normal"]))
            story.append(Spacer(1, 0.2 * cm))
        if i % 10 == 0:
            table_data = [["Name", "Email", "Score"]]
            for _ in range(6):
                table_data.append([fake.name(), fake.email(),
                    str(round(fake.pyfloat(min_value=0, max_value=100, right_digits=1), 1))])
            tbl = Table(table_data, colWidths=[5 * cm, 8 * cm, 3 * cm])
            tbl.setStyle(TableStyle([
                ("BACKGROUND",     (0, 0), (-1, 0), colors.HexColor("#4A90D9")),
                ("TEXTCOLOR",      (0, 0), (-1, 0), colors.white),
                ("FONTNAME",       (0, 0), (-1, 0), "Helvetica-Bold"),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.whitesmoke, colors.white]),
                ("GRID",           (0, 0), (-1, -1), 0.5, colors.grey),
                ("FONTSIZE",       (0, 0), (-1, -1), 9),
                ("PADDING",        (0, 0), (-1, -1), 4),
            ]))
            story.append(Spacer(1, 0.3 * cm))
            story.append(tbl)
        story.append(PageBreak())

    doc = SimpleDocTemplate(str(output), pagesize=A4)
    doc.build(story)


def gen_html(fake: Faker, sections: int, output: Path) -> None:
    import html as _hl

    parts = [
        "<!DOCTYPE html>", "<html lang='en'>", "<head>", "  <meta charset='UTF-8'>",
        f"  <title>{_hl.escape(fake.company())}</title>",
        "  <style>",
        "    body{font-family:Arial,sans-serif;max-width:960px;margin:0 auto;padding:2rem}",
        "    h1{color:#2c3e50}h2{color:#3498db;border-bottom:1px solid #eee}",
        "    table{border-collapse:collapse;width:100%;margin:1rem 0}",
        "    th{background:#3498db;color:white;padding:8px}",
        "    td{padding:6px 8px;border:1px solid #ddd}tr:nth-child(even){background:#f9f9f9}",
        "  </style>", "</head>", "<body>",
        f"<h1>{_hl.escape(fake.company())}</h1>",
        f"<p>{_hl.escape(fake.paragraph(nb_sentences=3))}</p>",
    ]
    for i in _progress(range(sections), label="Building HTML", total=sections):
        parts.append(f"<section id='s{i+1}'>")
        parts.append(f"  <h2>{_hl.escape(fake.catch_phrase())}</h2>")
        for _ in range(2):
            parts.append(f"  <p>{_hl.escape(fake.paragraph(nb_sentences=4))}</p>")
        parts.append("  <table><tr><th>Name</th><th>Email</th><th>Phone</th><th>Company</th></tr>")
        for _ in range(5):
            parts.append(
                f"  <tr><td>{_hl.escape(fake.name())}</td><td>{_hl.escape(fake.email())}</td>"
                f"<td>{_hl.escape(fake.phone_number())}</td><td>{_hl.escape(fake.company())}</td></tr>"
            )
        parts.append("  </table><ul>")
        for _ in range(4):
            parts.append(f"    <li>{_hl.escape(fake.sentence())}</li>")
        parts.append("  </ul></section>")
    parts.append("</body></html>")
    output.write_text("\n".join(parts), encoding="utf-8")


def gen_xml(fake: Faker, rows: int, output: Path) -> None:
    """Stdlib xml.etree -- no extra dependency."""
    root = ET.Element("records", count=str(rows), generated=datetime.now().isoformat())
    for i in _progress(range(rows), label="Building XML", total=rows):
        rec = ET.SubElement(root, "record", id=str(i + 1))
        for field, value in [
            ("name", fake.name()), ("email", fake.email()), ("phone", fake.phone_number()),
            ("city", fake.city()), ("country", fake.country()), ("company", fake.company()),
        ]:
            ET.SubElement(rec, field).text = value
    tree = ET.ElementTree(root)
    ET.indent(tree, space="  ")
    tree.write(str(output), encoding="utf-8", xml_declaration=True)


def gen_markdown(fake: Faker, sections: int, output: Path) -> None:
    """Stdlib only -- no extra dependency."""
    lines = [
        f"# {fake.catch_phrase()}\n\n",
        f"> Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')} | Sections: {sections}\n\n",
        "---\n\n",
    ]
    for i in _progress(range(sections), label="Building Markdown", total=sections):
        lines.append(f"## {i + 1}. {fake.sentence(nb_words=6).rstrip('.')}\n\n")
        lines.append(fake.paragraph(nb_sentences=4) + "\n\n")
        lines.append("| Name | Email | City | Score |\n")
        lines.append("|------|-------|------|-------|\n")
        for _ in range(5):
            lines.append(
                f"| {fake.name()} | {fake.email()} | {fake.city()} "
                f"| {round(fake.pyfloat(min_value=0, max_value=100, right_digits=2), 2)} |\n"
            )
        lines.append("\n```python\n")
        lines.append(f"# Sample snippet for section {i + 1}\n")
        lines.append(f"result = process(id={fake.pyint(min_value=1, max_value=9999)})\n")
        lines.append("```\n\n")
    output.write_text("".join(lines), encoding="utf-8")


def gen_image(fake: Faker, width: int, height: int, output: Path) -> None:
    """Requires Pillow (included in dependencies)."""
    from PIL import Image, ImageDraw

    img = Image.new("RGB", (width, height), color=(
        fake.random_int(20, 200), fake.random_int(20, 200), fake.random_int(20, 200),
    ))
    draw = ImageDraw.Draw(img)
    for _ in range(15):
        x1 = fake.random_int(0, max(1, width - 100))
        y1 = fake.random_int(0, max(1, height - 100))
        x2 = min(width,  x1 + fake.random_int(40, 200))
        y2 = min(height, y1 + fake.random_int(40, 150))
        draw.rectangle([x1, y1, x2, y2],
                       outline=(fake.random_int(100, 255), fake.random_int(100, 255), fake.random_int(100, 255)),
                       width=2)
    draw.text((10, 10), fake.catch_phrase()[:40], fill=(255, 255, 255))
    draw.text((10, height - 30), f"{width}x{height} | Test Image", fill=(200, 200, 200))
    fmt = "JPEG" if output.suffix.lower() in (".jpg", ".jpeg") else "PNG"
    img.save(str(output), fmt, **({"quality": 85} if fmt == "JPEG" else {}))


def gen_lorem(fake: Faker, words: int, output: Path) -> None:
    chunks, total = [], 0
    total_paras = max(1, words // 50)
    for _ in _progress(range(total_paras), label="Building text", total=total_paras):
        para = fake.paragraph(nb_sentences=5)
        chunks.append(para)
        total += len(para.split())
        if total >= words:
            break
    output.write_text("\n\n".join(chunks), encoding="utf-8")


def gen_zip(fake: Faker, output: Path) -> None:
    """Bundles CSV + JSON + Markdown + HTML. Stdlib zipfile -- no extra dep."""
    import tempfile

    with tempfile.TemporaryDirectory() as tmpdir:
        tmp = Path(tmpdir)
        gen_csv(fake, 200, tmp / "data.csv", "users")
        gen_json(fake, 200, tmp / "records.json", "users")
        gen_markdown(fake, 5, tmp / "readme.md")
        gen_html(fake, 3, tmp / "report.html")
        files = [tmp / "data.csv", tmp / "records.json", tmp / "readme.md", tmp / "report.html"]
        with zipfile.ZipFile(output, "w", compression=zipfile.ZIP_DEFLATED) as zf:
            for f in files:
                zf.write(f, arcname=f.name)
    click.echo(f"  bundled: {', '.join(f.name for f in files)}")


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

@click.group()
def cli() -> None:
    """Generate large test data files -- no venv needed, just `uv run generate.py`."""


def _common_options(fn):
    fn = click.option("--output", "-o", default=None, help="Output file path (default: auto-named).")(fn)
    fn = click.option("--seed", default=None, type=int, help="Integer seed for reproducible output.")(fn)
    fn = click.option("--locale", default="en_US", show_default=True,
                      help="Faker locale, e.g. fr_FR, vi_VN, ja_JP, de_DE.")(fn)
    return fn


def _schema_option(fn):
    return click.option(
        "--schema", type=click.Choice(list(SCHEMAS.keys())), default="users", show_default=True,
        help="Data schema / theme.",
    )(fn)


def _size_option(fn):
    return click.option(
        "--size", default=None, type=SIZE_CHOICES,
        help="Size preset: small/medium/large. Overridden by explicit --rows/--pages/--sections.",
    )(fn)


def _echo_done(out: Path) -> None:
    click.echo(click.style(f"  done: {out}  ({_file_size_label(out)})", fg="green"))


# -- JSON -------------------------------------------------------------------
@cli.command()
@click.option("--rows", default=None, type=int, help="Number of records (overrides --size).")
@_size_option
@_schema_option
@_common_options
def json(rows, size, schema, output, seed, locale):
    """Generate a JSON array file."""
    fake = _make_faker(locale, seed)
    n = _resolve(rows, "rows", size)
    out = Path(output) if output else Path(f"test_{schema}_{n}.json")
    click.echo(f"-> {n:,} records  schema={schema}  locale={locale}  -> {out}")
    gen_json(fake, n, out, schema)
    _echo_done(out)


# -- CSV --------------------------------------------------------------------
@cli.command()
@click.option("--rows", default=None, type=int, help="Number of rows (overrides --size).")
@_size_option
@_schema_option
@_common_options
def csv(rows, size, schema, output, seed, locale):
    """Generate a CSV file."""
    fake = _make_faker(locale, seed)
    n = _resolve(rows, "rows", size)
    out = Path(output) if output else Path(f"test_{schema}_{n}.csv")
    click.echo(f"-> {n:,} rows  schema={schema}  locale={locale}  -> {out}")
    gen_csv(fake, n, out, schema)
    _echo_done(out)


# -- XLSX -------------------------------------------------------------------
@cli.command()
@click.option("--rows", default=None, type=int, help="Number of rows (overrides --size).")
@_size_option
@_schema_option
@_common_options
def xlsx(rows, size, schema, output, seed, locale):
    """Generate an Excel (.xlsx) file."""
    fake = _make_faker(locale, seed)
    n = _resolve(rows, "rows", size)
    out = Path(output) if output else Path(f"test_{schema}_{n}.xlsx")
    click.echo(f"-> {n:,} rows  schema={schema}  locale={locale}  -> {out}")
    gen_xlsx(fake, n, out, schema)
    _echo_done(out)


# -- DOCX -------------------------------------------------------------------
@cli.command()
@click.option("--pages", default=None, type=int, help="Approximate page count (overrides --size).")
@_size_option
@_common_options
def docx(pages, size, output, seed, locale):
    """Generate a Word (.docx) document."""
    fake = _make_faker(locale, seed)
    n = _resolve(pages, "pages", size)
    out = Path(output) if output else Path(f"test_doc_{n}pages.docx")
    click.echo(f"-> ~{n} pages  locale={locale}  -> {out}")
    gen_docx(fake, n, out)
    _echo_done(out)


# -- PDF --------------------------------------------------------------------
@cli.command()
@click.option("--pages", default=None, type=int, help="Number of pages (overrides --size).")
@_size_option
@_common_options
def pdf(pages, size, output, seed, locale):
    """Generate a PDF file."""
    fake = _make_faker(locale, seed)
    n = _resolve(pages, "pages", size)
    out = Path(output) if output else Path(f"test_doc_{n}pages.pdf")
    click.echo(f"-> {n} pages  locale={locale}  -> {out}")
    gen_pdf(fake, n, out)
    _echo_done(out)


# -- HTML -------------------------------------------------------------------
@cli.command()
@click.option("--sections", default=None, type=int, help="Number of sections (overrides --size).")
@_size_option
@_common_options
def html(sections, size, output, seed, locale):
    """Generate an HTML page."""
    fake = _make_faker(locale, seed)
    n = _resolve(sections, "sections", size)
    out = Path(output) if output else Path(f"test_page_{n}sections.html")
    click.echo(f"-> {n} sections  locale={locale}  -> {out}")
    gen_html(fake, n, out)
    _echo_done(out)


# -- XML --------------------------------------------------------------------
@cli.command()
@click.option("--rows", default=None, type=int, help="Number of records (overrides --size).")
@_size_option
@_common_options
def xml(rows, size, output, seed, locale):
    """Generate an XML file (stdlib xml.etree, no extra dep)."""
    fake = _make_faker(locale, seed)
    n = _resolve(rows, "rows", size)
    out = Path(output) if output else Path(f"test_data_{n}.xml")
    click.echo(f"-> {n:,} records  locale={locale}  -> {out}")
    gen_xml(fake, n, out)
    _echo_done(out)


# -- Markdown ---------------------------------------------------------------
@cli.command()
@click.option("--sections", default=None, type=int, help="Number of sections (overrides --size).")
@_size_option
@_common_options
def md(sections, size, output, seed, locale):
    """Generate a Markdown (.md) file (stdlib only, no extra dep)."""
    fake = _make_faker(locale, seed)
    n = _resolve(sections, "sections", size)
    out = Path(output) if output else Path(f"test_doc_{n}sections.md")
    click.echo(f"-> {n} sections  locale={locale}  -> {out}")
    gen_markdown(fake, n, out)
    _echo_done(out)


# -- Image ------------------------------------------------------------------
@cli.command()
@click.option("--width",  default=None, type=int, help="Width in pixels (overrides --size).")
@click.option("--height", default=None, type=int, help="Height in pixels (overrides --size).")
@click.option("--format", "fmt", default="png", show_default=True,
              type=click.Choice(["png", "jpg"], case_sensitive=False),
              help="Image format.")
@_size_option
@_common_options
def image(width, height, fmt, size, output, seed, locale):
    """Generate a test image (PNG or JPG) using Pillow."""
    fake = _make_faker(locale, seed)
    w = _resolve(width,  "width",  size)
    h = _resolve(height, "height", size)
    out = Path(output) if output else Path(f"test_image_{w}x{h}.{fmt}")
    click.echo(f"-> {w}x{h}  format={fmt}  -> {out}")
    gen_image(fake, w, h, out)
    _echo_done(out)


# -- Lorem / plain text -----------------------------------------------------
@cli.command()
@click.option("--words", default=None, type=int, help="Approximate word count (overrides --size).")
@_size_option
@_common_options
def lorem(words, size, output, seed, locale):
    """Generate a plain-text file with realistic paragraph content."""
    fake = _make_faker(locale, seed)
    n = _resolve(words, "words", size)
    out = Path(output) if output else Path(f"test_text_{n}words.txt")
    click.echo(f"-> ~{n:,} words  locale={locale}  -> {out}")
    gen_lorem(fake, n, out)
    _echo_done(out)


# -- ZIP --------------------------------------------------------------------
@cli.command(name="zip")
@_common_options
def zip_cmd(output, seed, locale):
    """Generate a ZIP archive bundling CSV, JSON, Markdown, and HTML (stdlib zipfile)."""
    fake = _make_faker(locale, seed)
    out = Path(output) if output else Path("test_bundle.zip")
    click.echo(f"-> bundle  locale={locale}  -> {out}")
    gen_zip(fake, out)
    _echo_done(out)


# -- Batch ------------------------------------------------------------------
@cli.command()
@_size_option
@_common_options
def batch(size, output, seed, locale):
    """Generate one file of every format using the chosen size preset (default: medium)."""
    size = size or "medium"
    fake = _make_faker(locale, seed)
    out_dir = Path(output) if output else Path("test_batch")
    out_dir.mkdir(parents=True, exist_ok=True)
    p = SIZE_PRESETS[size]

    click.echo(f"-> Batch  size={size}  locale={locale}  -> {out_dir}/\n")

    def _run(label: str, fn, out: Path) -> None:
        click.echo(f"  [{label:6}] generating ...", nl=False)
        try:
            fn()
            click.echo(click.style(f" done  {out.name}  ({_file_size_label(out)})", fg="green"))
        except Exception as exc:
            click.echo(click.style(f" FAIL  {exc}", fg="red"))

    _run("json",  lambda: gen_json(     fake, p["rows"],               out_dir / f"data_{size}.json",  "users"),  out_dir / f"data_{size}.json")
    _run("csv",   lambda: gen_csv(      fake, p["rows"],               out_dir / f"data_{size}.csv",   "users"),  out_dir / f"data_{size}.csv")
    _run("xlsx",  lambda: gen_xlsx(     fake, p["rows"],               out_dir / f"data_{size}.xlsx",  "users"),  out_dir / f"data_{size}.xlsx")
    _run("docx",  lambda: gen_docx(     fake, p["pages"],              out_dir / f"doc_{size}.docx"),              out_dir / f"doc_{size}.docx")
    _run("pdf",   lambda: gen_pdf(      fake, p["pages"],              out_dir / f"doc_{size}.pdf"),               out_dir / f"doc_{size}.pdf")
    _run("html",  lambda: gen_html(     fake, p["sections"],           out_dir / f"page_{size}.html"),             out_dir / f"page_{size}.html")
    _run("xml",   lambda: gen_xml(      fake, p["rows"],               out_dir / f"data_{size}.xml"),              out_dir / f"data_{size}.xml")
    _run("md",    lambda: gen_markdown( fake, p["sections"],           out_dir / f"doc_{size}.md"),                out_dir / f"doc_{size}.md")
    _run("image", lambda: gen_image(    fake, p["width"], p["height"], out_dir / f"img_{size}.png"),               out_dir / f"img_{size}.png")
    _run("lorem", lambda: gen_lorem(    fake, p["words"],              out_dir / f"text_{size}.txt"),              out_dir / f"text_{size}.txt")
    _run("zip",   lambda: gen_zip(      fake,                          out_dir / f"bundle_{size}.zip"),            out_dir / f"bundle_{size}.zip")

    click.echo(click.style(f"\ndone: batch -> {out_dir.resolve()}", fg="green"))


if __name__ == "__main__":
    cli()
